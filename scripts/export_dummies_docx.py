#!/usr/bin/env python3
"""Export docs/HEXHAWK_FOR_DUMMIES.md to a Word .docx with embedded images.

This is a small purpose-built exporter for the Dummies guide. It preserves the
main reading structure, embeds local Markdown images, and keeps code blocks and
Markdown tables readable without relying on external tools like pandoc.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "HEXHAWK_FOR_DUMMIES.md"
OUTPUT = ROOT / "docs" / "HEXHAWK_FOR_DUMMIES.docx"
MAX_IMAGE_WIDTH = Inches(6.4)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_hyper_plain(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    """Add text with lightweight inline Markdown cleanup."""
    # Split backtick/code and bold markers enough for this manual.
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(31, 41, 55)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic


def add_code_block(doc: Document, code: list[str], language: str = "") -> None:
    if language:
        p = doc.add_paragraph()
        r = p.add_run(language)
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(100, 116, 139)
    p = doc.add_paragraph()
    p.style = "No Spacing"
    for idx, line in enumerate(code):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        if idx != len(code) - 1:
            run.add_break(WD_BREAK.LINE)
    # Lightly shade the paragraph via pPr.
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F5F9")
    p_pr.append(shd)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip().replace("`", "") for c in line.split("|")]


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    if len(rows) < 2:
        for row in rows:
            doc.add_paragraph(row)
        return
    header = parse_table_row(rows[0])
    body = [parse_table_row(r) for r in rows[2:]] if len(rows) > 2 else []
    cols = max(len(header), *(len(r) for r in body)) if body else len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    for i in range(cols):
        cell = table.rows[0].cells[i]
        cell.text = header[i] if i < len(header) else ""
        set_cell_shading(cell, "E2E8F0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row in body:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].text = row[i] if i < len(row) else ""
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def add_image(doc: Document, alt: str, rel: str) -> None:
    image_path = (SOURCE.parent / rel).resolve()
    p = doc.add_paragraph()
    if not image_path.exists():
        r = p.add_run(f"[Missing image: {rel}] {alt}")
        r.bold = True
        r.font.color.rgb = RGBColor(185, 28, 28)
        return
    p.alignment = 1
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=MAX_IMAGE_WIDTH)
    except Exception as exc:
        p.add_run(f"[Image could not be embedded: {rel}; {exc}]")
        return
    alt_p = doc.add_paragraph()
    alt_p.alignment = 1
    r = alt_p.add_run(f"Alt text: {alt}")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(71, 85, 105)


def build_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 24), ("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 12)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)

    doc.core_properties.title = "HexHawk for Dummies"
    doc.core_properties.subject = "Beginner guide with embedded screenshots"
    doc.core_properties.author = "HexHawk / Hermes Agent"

    lines = text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code: list[str] = []
    table_rows: list[str] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code = []
            else:
                add_code_block(doc, code, code_lang)
                in_code = False
                code_lang = ""
                code = []
            i += 1
            continue

        if in_code:
            code.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows.append(line)
            i += 1
            continue
        else:
            flush_table()

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph().add_run("—" * 20).font.color.rgb = RGBColor(148, 163, 184)
            i += 1
            continue

        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            add_image(doc, img.group(1), img.group(2))
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_hyper_plain(p, stripped[2:])
        elif re.match(r"\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_hyper_plain(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("Caption:"):
            p = doc.add_paragraph()
            r = p.add_run("Caption: ")
            r.bold = True
            r.italic = True
            add_hyper_plain(p, stripped[len("Caption:"):].strip(), italic=True)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(71, 85, 105)
        else:
            p = doc.add_paragraph()
            add_hyper_plain(p, stripped)
        i += 1

    flush_table()
    if in_code:
        add_code_block(doc, code, code_lang)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
